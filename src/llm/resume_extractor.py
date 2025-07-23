import os
import json
import time
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging
import PyPDF2
import docx
import pandas as pd
import re
from dotenv import load_dotenv
from openai import OpenAI
import asyncio
from collections import defaultdict, deque
import subprocess
import platform
from pathlib import Path
import base64

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
import numpy as np
import cv2

from unicodedata import category
import unicodedata

from google.cloud import vision
from google.oauth2 import service_account
import tempfile

# Configure logger
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

load_dotenv()

class ResumeExtractor:
    """Enhanced resume extractor with OpenAI Vision API for scanned documents"""
    
    def __init__(self):
        """Initialize the multi-API resume extractor."""

        # Google Cloud Vision API configuration
        self.google_credentials_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
        self.google_project_id = os.getenv('GOOGLE_CLOUD_PROJECT_ID')
        self.google_credentials_json = os.getenv('GOOGLE_APPLICATION_CREDENTIALS_JSON')

        # API configurations
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("OPENAI_API_KEY is not set")

        self.vision_client = None
        self.temp_file_path = None

        try:
            if self.google_credentials_path and os.path.exists(self.google_credentials_path):
                # Use the credentials file if it exists
                credentials = service_account.Credentials.from_service_account_file(
                    self.google_credentials_path
                )
                logger.info("Using Google Cloud credentials from file")
                
            elif self.google_credentials_json:
                # Parse and validate JSON first
                try:
                    credentials_dict = json.loads(self.google_credentials_json)
                except json.JSONDecodeError as e:
                    raise ValueError(f"Invalid JSON in GOOGLE_APPLICATION_CREDENTIALS_JSON: {e}")
                
                # Create credentials directly from dict (more efficient)
                credentials = service_account.Credentials.from_service_account_info(
                    credentials_dict
                )
                logger.info("Using Google Cloud credentials from environment variable")
                
            else:
                raise ValueError("No valid Google Cloud credentials provided. Set either GOOGLE_APPLICATION_CREDENTIALS or GOOGLE_APPLICATION_CREDENTIALS_JSON")

            self.vision_client = vision.ImageAnnotatorClient(credentials=credentials)
            logger.info("Google Cloud Vision API initialized successfully")

        except Exception as e:
            logger.error(f"Failed to initialize Google Cloud Vision API: {e}")
            raise
        
        # Available API providers and models
        self.api_providers = {
            'openai': {
                'vision_models': [('gpt-4o', 8000), ('gpt-4o-mini', 4000)],
                'text_models': [('gpt-4.1-mini', 4000)],
                'api_key': self.openai_api_key,
            },
            'google_vision': {
                'enabled': self.vision_client is not None,
                'client': self.vision_client,
            }
        }

        self.extraction_prompt = """
        You are an expert resume parser. Extract the following information from the resume provided.
        Return the data in JSON format with the exact structure shown below:

        {
            "personal_info": {
                "name": "Full name of the candidate",
                "email": "Email address",
                "phone": "Full phone number including country code if available",
                "location": "City, State/Country",
                "linkedin": "LinkedIn profile URL",
                #"portfolio": "Portfolio/website URL",
            },
            "summary": "Professional summary or objective statement",
            "skills": [
                "List of technical and soft skills",
            ],
            "experience": [
                {
                "title": "Job title",
                "company": "Company name",
                "location": "Job location",
                "duration": "Start date - End date",
                }
            ],
            "total_experience": "Calculate the total work experience from a list of non-overlapping date ranges. Treat 'Present' as the current date. Do not round or merge separate time periods. Count each month precisely, and sum them exactly. Output the total duration in years and months, e.g., '11 months' or '2 years 3 months'.",
            "education": [
                {
                "degree": "Degree name",
                "institution": "Institution name",
                "location": "Institution location",
                "year": "Graduation year or duration",
                "gpa": "GPA if mentioned",
                }
            ],
            "certifications": [
                {
                "name": "Certification name",
                "issuer": "Issuing organization",
                "date": "Date obtained",
                "expiry": "Expiry date if any",
                }
            ],
            "projects": [
                {
                "name": "Project name",
                "description": "Project description",
                "technologies": ["Technologies used"],
                "duration": "Project duration",
                }
            ],
            "languages": [
                {
                "language": "Language name",
                "proficiency": "Proficiency level",
                }
            ],
            "awards": [
                "List of awards and achievements",
            ]
            }

    **Extraction Rules:**
    - Extract all entries in each category (multiple jobs, degrees, projects, etc.).
    - Use empty strings or empty arrays for missing or not found values.
    - Include all emails and phone numbers found.
    - Include country code in phone numbers if present.
    - When extracting phone numbers, ignore visual dividers such as `|` or `/` and do not treat them as digits.
    - Only digits, spaces, parentheses, hyphens, and `+` are valid in phone numbers.
    - If the phone number ends in an extra digit due to a pipe or special character, remove it.
    - Include all skill types: Technical, Soft, Tools, Platforms, and Domain-Specific.
    - Only include skill from skill section not from other section.
    - Preserve bullet points in description fields where applicable.
    - Accept ALL-CAPS or spaced section headers (e.g., "P R O J E C T S") as valid dividers.
    - For sections titled "Projects", "Academic Projects", etc.:
    - Parse only into the projects array, not experience.
    - For the experience array, only include an entry if:
    - It names a company or organization, and
    - It includes a job title (e.g., Intern, Analyst, Developer).
    - Under experience, parse lines like X - Y as:
    - title = X, company = Y
    - If "Fresher" is mentioned or there's no employment history:
    - Keep the experience array empty.
    - Use contextual cues (e.g., "developed", "collaborated", "built") to recognize project entries.
    """

        self.ocr_methods = ['google_vision']
        
        # Current provider and model indices
        self.current_provider = 'openai'
        self.current_vision_model_index = 0
        self.current_text_model_index = 0
        
        # Usage tracking
        self.usage_file = 'resume_api_usage_tracking.json'
        self.usage_data = self._load_usage_data()
        
        # Rate limiting - aligned with main.py settings
        self.daily_limits = {
            'openai': 9500  # RPM limit
        }
        
        # Initialize OpenAI client
        self.openai_client = OpenAI(api_key=self.openai_api_key) if self.openai_api_key else None
        
        print(f"Initialized with provider: {self.current_provider}")
        print(f"Current vision model: {self.get_current_vision_model()}")
        print(f"Current text model: {self.get_current_text_model()}")

    def __del__(self):
        # Clean up temp file if it was created
        if hasattr(self, 'temp_file_path') and self.temp_file_path and os.path.exists(self.temp_file_path):
            try:
                os.unlink(self.temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to clean up temp file: {e}")

    def load_prompts(self):
        prompts = {}
        PROMPT_FOLDER = "src/prompts"
        for fname in os.listdir(PROMPT_FOLDER):
            if fname.endswith(".txt"):
                key = fname.replace(".txt", "")
                with open(os.path.join(PROMPT_FOLDER, fname), "r", encoding="utf-8") as f:
                    prompts[key] = f.read()
        return prompts

    def get_current_vision_model(self) -> str:
        """Get the current vision model being used."""
        provider_config = self.api_providers[self.current_provider]
        return provider_config['vision_models'][self.current_vision_model_index][0]

    def get_current_text_model(self) -> str:
        """Get the current text model being used."""
        provider_config = self.api_providers[self.current_provider]
        return provider_config['text_models'][self.current_text_model_index][0]
    
    def _extract_text_with_google_vision(self, file_path: str, max_pages: int = 10) -> str:
        """Extract text from PDF using Google Cloud Vision API OCR"""
        if not self.vision_client:
            raise Exception("Google Cloud Vision API not initialized")
        
        try:
            images = self._pdf_to_images_for_ocr(file_path, max_pages)
            if not images:
                raise Exception("Could not convert PDF to images for OCR")
            
            extracted_text = ""
            
            for i, image_data in enumerate(images):
                try:
                    image = vision.Image(content=image_data)
                    
                    # Use document text detection for better layout analysis
                    response = self.vision_client.document_text_detection(image=image)
                    
                    if response.error.message:
                        raise Exception(f"Google Vision API error: {response.error.message}")
                    
                    # Process with layout information
                    page_text = self._process_document_response(response)
                    extracted_text += f"\n--- Page {i+1} ---\n{page_text}\n"
                    
                    logger.info(f"Extracted text from page {i+1} with layout analysis")
                    
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1}: {e}")
                    continue
            
            return extracted_text.strip()
            
        except Exception as e:
            logger.error(f"Enhanced OCR failed: {e}")
            raise

    def _process_document_response(self, response) -> str:
        """Process Google Vision document response with layout preservation"""
        if not response.full_text_annotation:
            return ""
        
        # Group text by blocks and paragraphs
        text_blocks = []
        
        for page in response.full_text_annotation.pages:
            for block in page.blocks:
                block_text = []
                block_confidence = block.confidence
                
                for paragraph in block.paragraphs:
                    para_words = []
                    for word in paragraph.words:
                        word_text = ''.join([symbol.text for symbol in word.symbols])
                        para_words.append(word_text)
                    
                    if para_words:
                        paragraph_text = ' '.join(para_words)
                        block_text.append(paragraph_text)
                
                if block_text:
                    # Join paragraphs in block with newlines
                    full_block_text = '\n'.join(block_text)
                    text_blocks.append({
                        'text': full_block_text,
                        'confidence': block_confidence,
                        'bbox': self._get_block_bbox(block)
                    })
        
        # Sort blocks by position (top to bottom, left to right)
        text_blocks.sort(key=lambda x: (x['bbox']['top'], x['bbox']['left']))
        
        # Join blocks with appropriate spacing
        result_text = []
        for block in text_blocks:
            if block['confidence'] > 0.5:  # Filter low-confidence blocks
                result_text.append(block['text'])
        
        return '\n\n'.join(result_text)
    
    def _get_block_bbox(self, block) -> Dict[str, int]:
        """Get bounding box coordinates for a text block"""
        vertices = block.bounding_box.vertices
        
        if not vertices:
            return {'left': 0, 'top': 0, 'right': 0, 'bottom': 0}
        
        left = min(v.x for v in vertices)
        right = max(v.x for v in vertices)
        top = min(v.y for v in vertices)
        bottom = max(v.y for v in vertices)
        
        return {
            'left': left,
            'top': top, 
            'right': right,
            'bottom': bottom
        }
    
    def _pdf_to_images_for_ocr(self, file_path: str, max_pages: int = 10) -> List[bytes]:
        """Convert PDF pages to image bytes for OCR processing"""
        images = []
        
        try:
            doc = fitz.open(file_path)
            
            # Limit number of pages to avoid processing too many pages
            num_pages = min(len(doc), max_pages)
            
            for page_num in range(num_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Convert page to high-resolution image for better OCR
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                    img_data = pix.tobytes("png")
                    
                    images.append(img_data)
                    logger.info(f"Converted page {page_num + 1} to image for OCR")
                    
                except Exception as e:
                    logger.warning(f"Failed to convert page {page_num} to image: {e}")
                    continue
            
            doc.close()
            
        except Exception as e:
            logger.error(f"Error converting PDF to images for OCR: {e}")
            raise
        
        return images
    
    def _extract_with_ocr_fallback(self, file_path: str, filename: str, max_retries: int = 3) -> Dict[str, Any]:
        """
        Extract text using OCR methods with fallback chain:
        1. Google Cloud Vision API (if available)
        2. Tesseract OCR (local)
        3. OpenAI Vision API (as last resort)
        """
        ocr_text = ""
        extraction_method = ""
        
        # Try OCR methods in order of preference
        for method in self.ocr_methods:
            try:
                if method == 'google_vision' and self.vision_client:
                    logger.info(f"Attempting Google Cloud Vision OCR for {filename}")
                    ocr_text = self._extract_text_with_google_vision(file_path)
                    extraction_method = "google_vision_ocr"
                    
                """elif method == 'tesseract':
                    logger.info(f"Attempting Tesseract OCR for {filename}")
                    ocr_text = self._extract_text_with_tesseract(file_path)
                    extraction_method = "tesseract_ocr"
                    
                elif method == 'openai_vision' and self.openai_client:
                    logger.info(f"Attempting OpenAI Vision API for {filename}")
                    model = self.get_current_vision_model()
                    return self._extract_with_vision_api(file_path, filename, model, max_retries)"""
                
                logger.info(f"extarcted data {ocr_text}")
                # If we got good text, proceed with processing
                if ocr_text and self._validate_extracted_text(ocr_text):
                    logger.info(f"Successfully extracted text using {method} for {filename}")
                    break
                else:
                    logger.warning(f"OCR method {method} produced low-quality text for {filename}")
                    
            except Exception as e:
                logger.warning(f"OCR method {method} failed for {filename}: {e}")
                continue
        
        # If no OCR method worked, return error
        if not ocr_text or not self._validate_extracted_text(ocr_text):
            return {
                "filename": filename,
                "error": "All OCR methods failed to extract readable text",
                "success": False
            }
        
        # Process the extracted text with AI for structured data extraction
        model = self.get_current_text_model()
        logger.info(f"Processing OCR text with {model} for {filename}")
        
        for attempt in range(max_retries):
            try:
                response_text = self._make_api_call(ocr_text, self.current_provider, model, is_image=False)
                
                if response_text:
                    # Clean the response text to extract JSON
                    json_data = self._clean_json_response(response_text)
                    
                    # Add metadata
                    json_data["filename"] = filename
                    json_data["success"] = True
                    json_data["extraction_method"] = extraction_method
                    json_data["ocr_text_length"] = len(ocr_text)
                    json_data["provider"] = self.current_provider
                    json_data["model"] = model
                    json_data["extraction_timestamp"] = datetime.now().isoformat()
                    
                    # Post-process and validate data
                    #json_data = self._post_process_data(json_data)
                    
                    return json_data
                    
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error on attempt {attempt + 1}: {str(e)}")
                if attempt == max_retries - 1:
                    return {
                        "filename": filename,
                        "error": f"Failed to parse AI response as JSON after {max_retries} attempts: {str(e)}",
                        "raw_response": response_text[:500] if response_text else "No response",
                        "success": False
                    }
                time.sleep(1 * (attempt + 1))
                
            except Exception as e:
                logger.error(f"Error on attempt {attempt + 1}: {str(e)}")
                if attempt == max_retries - 1:
                    return {
                        "filename": filename,
                        "error": f"Failed to extract resume data after {max_retries} attempts: {str(e)}",
                        "success": False
                    }
                time.sleep(1 * (attempt + 1))
        
        return {
            "filename": filename,
            "error": f"Failed to extract resume data after {max_retries} attempts.",
            "success": False
        }


    def _make_openai_text_api_call(self, text_content: str, model: str) -> Optional[str]:
        """Make API call to OpenAI with text content (for text-based PDFs)"""
        prompts = self.load_prompts()
        if not self.openai_client:
            logger.error("OpenAI client not initialized - check API key")
            return None
            
        try:
            #safe_text = self.clean_unicode(text_content)
            all_prompts = "\n\n".join(prompts.values())
            full_prompt = self.extraction_prompt + "\n\nResume text:\n" + text_content
            
            # Determine max_tokens based on model
            if model == 'gpt-4.1-mini':
                max_output_tokens = 4000
            elif model == 'gpt-4.1-mini':
                max_output_tokens = 4000
            else:  # gpt-4o
                max_output_tokens = 4000
            
            completion = self.openai_client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": """You are a precise resume parser. Your job is to extract information EXACTLY as written in the resume without any modifications, cleaning, or truncation. 

                        CRITICAL RULES:
                        - Extract complete information - never truncate names, emails, or any other fields
                        - If a section is not present, return empty string or empty array
                        - Extract only from relevant sections - don't mix information from different sections
                        - Return only valid JSON without any markdown formatting
                        - Do not infer or generate information not explicitly stated
                        - ENSURE the JSON is complete and properly closed with all brackets and braces"""
                    },
                    {
                        "role": "user", 
                        "content": full_prompt
                    }
                ],
                max_tokens=max_output_tokens,
                temperature=0,  # Keep deterministic
            )

            # Update usage tracking
            self._update_api_usage('openai', model, 1)
            logger.info(completion.choices[0].message.content)
            return completion.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error making OpenAI text API call: {str(e)}")
            return None


    def _pdf_to_images(self, file_path: str, max_pages: int = 10) -> List[str]:
        """Convert PDF pages to base64 encoded images for Vision API"""
        images = []
        
        try:
            doc = fitz.open(file_path)
            
            # Limit number of pages to avoid token limits
            num_pages = min(len(doc), max_pages)
            
            for page_num in range(num_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Convert page to high-resolution image
                    mat = fitz.Matrix(3, 3)  # 3x zoom for good quality
                    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                    img_data = pix.tobytes("png")
                    
                    # Convert to base64
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    images.append(img_base64)
                    
                    logger.info(f"Converted page {page_num + 1} to image")
                    
                except Exception as e:
                    logger.warning(f"Failed to convert page {page_num} to image: {e}")
                    continue
            
            doc.close()
            
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            raise
        
        return images

    def _load_usage_data(self) -> Dict:
        """Load API usage data from file."""
        try:
            if os.path.exists(self.usage_file):
                with open(self.usage_file, 'r') as f:
                    return json.load(f)
            else:
                return {}
        except Exception as e:
            logger.error(f"Error loading usage data: {e}")
            return {}

    def _save_usage_data(self):
        """Save API usage data to file."""
        try:
            with open(self.usage_file, 'w') as f:
                json.dump(self.usage_data, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving usage data: {e}")

    def _get_today_key(self) -> str:
        """Get today's date as a string key."""
        return datetime.now().strftime('%Y-%m-%d')

    def _update_api_usage(self, provider: str, model: str, count: int = 1):
        """Update API usage for a provider and model."""
        today = self._get_today_key()
        
        if provider not in self.usage_data:
            self.usage_data[provider] = {}
        
        if model not in self.usage_data[provider]:
            self.usage_data[provider][model] = {}
        
        if today not in self.usage_data[provider][model]:
            self.usage_data[provider][model][today] = 0
        
        self.usage_data[provider][model][today] += count
        self._save_usage_data()

    def _get_today_usage(self, provider: str, model: str = None) -> int:
        """Get today's API usage for a provider (and optionally specific model)."""
        today = self._get_today_key()
        
        if model:
            return self.usage_data.get(provider, {}).get(model, {}).get(today, 0)
        else:
            # Get total usage for provider across all models
            total = 0
            provider_data = self.usage_data.get(provider, {})
            for model_data in provider_data.values():
                total += model_data.get(today, 0)
            return total

    def get_current_model(self) -> str:
        """Get the current model being used."""
        provider_config = self.api_providers[self.current_provider]
        return provider_config['models'][self.current_model_index][0]

    def _make_api_call(self, content, provider: str, model: str, is_image: bool = False) -> Optional[str]:
        """Make API call to extract resume data."""
        if provider == 'openai':
            if is_image:
                return self._make_openai_vision_api_call(content, model)
            else:
                return self._make_openai_text_api_call(content, model)
        else:
            logger.error(f"Unsupported provider: {provider}")
            return None

    def extract_from_file(self, file_path: str, filename: str, max_retries: int = 3) -> Dict[str, Any]:
        """
        Extract resume data from a file using appropriate model based on content type
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Check if we've hit rate limits
            if self._check_rate_limits():
                return {
                    "filename": filename,
                    "error": "Rate limit exceeded. Please try again later.",
                    "success": False
                }
            
            # Handle PDF files - always use OCR with Google Vision
            if file_extension == '.pdf':
                logger.info(f"Processing PDF with Google Vision OCR: {filename}")
                return self._extract_with_ocr_fallback(file_path, filename, max_retries)
            
            # Handle other file types with text model (they are all text-based)
            else:
                text_content = self._extract_text_from_file(file_path)
                
                if not text_content.strip():
                    return {
                        "filename": filename,
                        "error": "No text content could be extracted from the file",
                        "success": False
                    }
                
                # Use text model for non-PDF files
                model = self.get_current_text_model()
                logger.info(f"Using text model {model} for {filename}")
                return self._extract_with_text_api(text_content, filename, model, max_retries)
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return {
                "filename": filename,
                "error": f"Error processing file: {str(e)}",
                "success": False
            }

    def _extract_from_pdf_intelligent(self, file_path: str, filename: str, max_retries: int) -> Dict[str, Any]:
        """
        Always use Google Vision OCR for PDFs
        """
        try:
            logger.info(f"Using Google Vision OCR for PDF: {filename}")
            return self._extract_with_ocr_fallback(file_path, filename, max_retries)
        except Exception as e:
            logger.error(f"Error in PDF processing: {str(e)}")
            return {
                "filename": filename,
                "error": f"Error in PDF processing: {str(e)}",
                "success": False
            }
        
    def get_ocr_capabilities(self) -> Dict[str, bool]:
        """Get available OCR capabilities"""
        capabilities = {
            'google_vision': self.vision_client is not None,
            #'tesseract': self._check_tesseract_availability(),
            #'openai_vision': self.openai_client is not None
        }
        
        return capabilities

    def _extract_with_vision_api(self, file_path: str, filename: str, model: str, max_retries: int) -> Dict[str, Any]:
        """Extract resume data using Vision API for scanned/image-based PDFs"""
        try:
            # Convert PDF to images
            images = self._pdf_to_images(file_path)
            
            if not images:
                return {
                    "filename": filename,
                    "error": "Could not convert PDF to images",
                    "success": False
                }
            
            # Make API call with retries
            for attempt in range(max_retries):
                try:
                    response_text = self._make_api_call(images, self.current_provider, model, is_image=True)
                    
                    if response_text:
                        # Clean the response text to extract JSON
                        json_data = self._clean_json_response(response_text)
                        
                        # Add metadata
                        json_data["filename"] = filename
                        json_data["success"] = True
                        json_data["extraction_method"] = "vision_api"
                        json_data["pages_processed"] = len(images)
                        json_data["provider"] = self.current_provider
                        json_data["model"] = model
                        json_data["extraction_timestamp"] = datetime.now().isoformat()
                        
                        # Post-process and validate data
                        json_data = self._post_process_data(json_data)
                        
                        return json_data
                        
                except json.JSONDecodeError as e:
                    logger.error(f"JSON decode error on attempt {attempt + 1}: {str(e)}")
                    if attempt == max_retries - 1:
                        return {
                            "filename": filename,
                            "error": f"Failed to parse Vision API response as JSON after {max_retries} attempts: {str(e)}",
                            "raw_response": response_text[:500] if response_text else "No response",
                            "success": False
                        }
                    time.sleep(1 * (attempt + 1))
                    
                except Exception as e:
                    logger.error(f"Vision API error on attempt {attempt + 1}: {str(e)}")
                    if attempt == max_retries - 1:
                        return {
                            "filename": filename,
                            "error": f"Failed to extract resume data using Vision API after {max_retries} attempts: {str(e)}",
                            "success": False
                        }
                    time.sleep(1 * (attempt + 1))
            
            return {
                "filename": filename,
                "error": f"Failed to extract resume data after {max_retries} attempts.",
                "success": False
            }
            
        except Exception as e:
            logger.error(f"Error in vision API extraction: {str(e)}")
            return {
                "filename": filename,
                "error": f"Error in vision API processing: {str(e)}",
                "success": False
            }
    
    def _extract_with_text_api(self, text_content: str, filename: str, model: str, max_retries: int) -> Dict[str, Any]:
        """Extract resume data using text-based API"""
        for attempt in range(max_retries):
            try:
                response_text = self._make_api_call(text_content, self.current_provider, model, is_image=False)
                
                if response_text:
                    # Clean the response text to extract JSON
                    json_data = self._clean_json_response(response_text)
                    
                    # Add metadata
                    json_data["filename"] = filename
                    json_data["success"] = True
                    json_data["text_length"] = len(text_content)
                    json_data["extraction_method"] = "text_api"
                    json_data["provider"] = self.current_provider
                    json_data["model"] = model
                    json_data["extraction_timestamp"] = datetime.now().isoformat()
                    
                    # Post-process and validate data
                    json_data = self._post_process_data(json_data)
                    
                    return json_data
                    
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error on attempt {attempt + 1}: {str(e)}")
                if attempt == max_retries - 1:
                    return {
                        "filename": filename,
                        "error": f"Failed to parse AI response as JSON after {max_retries} attempts: {str(e)}",
                        "raw_response": response_text[:500] if response_text else "No response",
                        "success": False
                    }
                time.sleep(1 * (attempt + 1))
                
            except Exception as e:
                logger.error(f"Error on attempt {attempt + 1}: {str(e)}")
                if attempt == max_retries - 1:
                    return {
                        "filename": filename,
                        "error": f"Failed to extract resume data after {max_retries} attempts: {str(e)}",
                        "success": False
                    }
                time.sleep(1 * (attempt + 1))
        
        return {
            "filename": filename,
            "error": f"Failed to extract resume data after {max_retries} attempts.",
            "success": False
        }

    def _check_rate_limits(self) -> bool:
        """Check if we've exceeded rate limits."""
        today_usage = self._get_today_usage(self.current_provider)
        daily_limit = self.daily_limits.get(self.current_provider, 1000)
        return today_usage >= daily_limit

    def _validate_extracted_text(self, text: str) -> bool:
        """Validate if extracted text is of good quality for resume parsing"""
        if not text or len(text.strip()) < 100:
            return False
        
        # Check for common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'email', 'phone', 'work',
            'university', 'college', 'project', 'internship', 'job', 'career',
            'resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary'
        ]
        
        text_lower = text.lower()
        found_indicators = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        # Check for email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        has_email = bool(re.search(email_pattern, text))
        
        # Check for phone pattern
        phone_pattern = r'[\+]?[\d\s\-\(\)]{10,}'
        has_phone = bool(re.search(phone_pattern, text))
        
        # Check for meaningful word density
        words = text.split()
        meaningful_words = [word for word in words if len(word) > 2 and word.isalpha()]
        word_density = len(meaningful_words) / len(words) if words else 0
        
        return (found_indicators >= 2 and word_density > 0.4) or has_email or has_phone

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from various file formats"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                return self._extract_from_doc(file_path)
            elif file_extension in ['.xls', '.xlsx']:
                return self._extract_from_excel(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension} file: {str(e)}")
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")


    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sheet_text = df.fillna('').to_string(index=False)
                    text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
                except Exception as e:
                    logger.warning(f"Error reading sheet {sheet_name}: {e}")
                    continue
        except Exception as e:
            raise Exception(f"Error reading Excel: {str(e)}")
        
        return text.strip()

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        text = ""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if not text:
            raise Exception("Could not read TXT file with any supported encoding")
        
        return text.strip()

    def clean_unicode(self, text: str) -> str:
        """Enhanced unicode cleaning for better text extraction"""
        if not text:
            return ""
        
        # Normalize unicode characters
        text = unicodedata.normalize("NFKD", text)
        
        # Remove control characters but keep essential whitespace
        cleaned = ''.join(c for c in text if category(c)[0] != "C" or c in '\n\r\t ')
        
        # Fix common OCR errors - but be careful with phone number context
        ocr_fixes = {
            'rn': 'm',  # rn combination often misread as m
            '€': 'e',  # Euro symbol instead of e
            '©': 'c',  # Copyright symbol instead of c
            '®': 'r',  # Registered symbol instead of r
            '™': 'tm', # Trademark symbol
            '"': '"',  # Smart quotes
            '"': '"',  # Smart quotes
            ''': "'",  # Smart apostrophes
            ''': "'",  # Smart apostrophes
            '–': '-',  # En dash to hyphen
            '—': '-',  # Em dash to hyphen
            '…': '...',  # Ellipsis
        }
        
        # Apply OCR fixes but avoid phone number context
        for wrong, correct in ocr_fixes.items():
            cleaned = cleaned.replace(wrong, correct)
        
        # Remove excessive whitespace
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # Remove lines with only punctuation or numbers (likely OCR artifacts)
        lines = cleaned.split('\n')
        filtered_lines = []
        for line in lines:
            line = line.strip()
            if line and not re.match(r'^[^\w\s]*$', line):  # Keep lines with at least some alphanumeric
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)

    def _clean_phone_number(self, phone: str) -> str:
        """Clean and validate phone number with better handling of OCR artifacts"""
        if not phone:
            return ""
        
        # First, remove common OCR artifacts that appear at the end of phone numbers
        phone = phone.strip()
        
        # Remove trailing pipes and other common OCR artifacts
        phone = re.sub(r'[|/\\]+$', '', phone)
        
        # Remove any characters that are clearly not part of a phone number
        # Keep only digits, +, -, (, ), and spaces
        phone = re.sub(r'[^\d\+\-\(\)\s]', '', phone)
        
        # Normalize spaces
        phone = re.sub(r'\s+', ' ', phone).strip()
        
        # Validate phone number (should have at least 10 digits)
        digits_only = re.sub(r'\D', '', phone)
        
        # Check if we have a reasonable number of digits (10-15 is typical for international numbers)
        if 10 <= len(digits_only) <= 15:
            return phone
        else:
            return ""
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from legacy DOC file using multiple approaches"""
        text = ""
        
        # Method 1: Try using python-docx2txt (if available)
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
        except ImportError:
            logger.info("docx2txt not available, trying alternative methods")
        except Exception as e:
            logger.warning(f"docx2txt failed: {e}")
        
        # Method 2: Try using antiword (Linux/Mac command line tool)
        try:
            if platform.system() in ['Linux', 'Darwin']:  # Linux or Mac
                result = subprocess.run(
                    ['antiword', file_path], 
                    capture_output=True, 
                    text=True, 
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
            logger.warning(f"antiword failed: {e}")
        
        # Method 3: Try using catdoc (Linux/Mac command line tool)
        try:
            if platform.system() in ['Linux', 'Darwin']:  # Linux or Mac
                result = subprocess.run(
                    ['catdoc', file_path], 
                    capture_output=True, 
                    text=True, 
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
            logger.warning(f"catdoc failed: {e}")
        
        # Method 4: Try using LibreOffice/OpenOffice conversion (if available)
        try:
            temp_dir = Path(file_path).parent / "temp_conversion"
            temp_dir.mkdir(exist_ok=True)
            
            # Try LibreOffice first
            commands_to_try = [
                ['libreoffice', '--headless', '--convert-to', 'txt', '--outdir', str(temp_dir), file_path],
                ['soffice', '--headless', '--convert-to', 'txt', '--outdir', str(temp_dir), file_path]
            ]
            
            for cmd in commands_to_try:
                try:
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                    if result.returncode == 0:
                        # Find the converted txt file
                        txt_filename = Path(file_path).stem + '.txt'
                        txt_path = temp_dir / txt_filename
                        
                        if txt_path.exists():
                            with open(txt_path, 'r', encoding='utf-8') as f:
                                converted_text = f.read()
                            
                            # Clean up temp file
                            txt_path.unlink()
                            temp_dir.rmdir()
                            
                            if converted_text.strip():
                                return converted_text.strip()
                    
                except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
                    logger.warning(f"LibreOffice conversion failed: {e}")
                    continue
            
            # Clean up temp directory if it exists
            if temp_dir.exists():
                try:
                    temp_dir.rmdir()
                except:
                    pass
                    
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {e}")
        
        # Method 5: Try reading as binary and extract readable text (fallback)
        try:
            with open(file_path, 'rb') as f:
                binary_data = f.read()
            
            # Try to decode and extract readable text
            try:
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        decoded = binary_data.decode(encoding, errors='ignore')
                        # Extract printable ASCII characters
                        import string
                        printable_text = ''.join(char for char in decoded if char in string.printable)
                        
                        # Clean up the text - remove excessive whitespace and non-text elements
                        lines = []
                        for line in printable_text.split('\n'):
                            line = line.strip()
                            if len(line) > 2 and not line.startswith('\n'):
                                lines.append(line)
                        
                        cleaned_text = '\n'.join(lines)
                        
                        # If we have reasonable amount of text, return it
                        if len(cleaned_text) > 50:  # Minimum threshold
                            return cleaned_text
                            
                    except UnicodeDecodeError:
                        continue
                        
            except Exception as e:
                logger.warning(f"Binary text extraction failed: {e}")
                
        except Exception as e:
            logger.warning(f"Binary file reading failed: {e}")
        
        # If all methods fail, raise an exception with helpful information
        raise Exception(
            f"Could not extract text from .doc file. "
            f"Please consider converting the file to .docx format or install additional tools:\n"
            f"- Install docx2txt: pip install docx2txt\n"
            f"- Install antiword (Linux/Mac): sudo apt-get install antiword or brew install antiword\n"
            f"- Install catdoc (Linux/Mac): sudo apt-get install catdoc or brew install catdoc\n"
            f"- Install LibreOffice for conversion support"
        )

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with better error handling"""
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
        except Exception as e:
            # If it's a .docx file but fails, it might actually be a .doc file with wrong extension
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.docx':
                logger.warning(f"DOCX extraction failed, trying DOC methods: {e}")
                try:
                    return self._extract_from_doc(file_path)
                except Exception as doc_error:
                    raise Exception(f"Error reading DOCX (tried DOC methods too): {str(e)} | DOC attempt: {str(doc_error)}")
            else:
                raise Exception(f"Error reading DOCX: {str(e)}")
        
        return text.strip()

    def _clean_json_response(self, response_text: str) -> dict:
        """Clean and safely parse the AI response into valid JSON"""
        if not response_text:
            return {}

        # Step 1: Remove markdown code block markers
        response_text = re.sub(r'```json\s*', '', response_text, flags=re.IGNORECASE)
        response_text = re.sub(r'```\s*$', '', response_text, flags=re.MULTILINE)
        response_text = response_text.strip()

        # Step 2: Try to parse as single JSON object first
        try:
            # Clean the text before parsing
            cleaned_text = self._sanitize_json_text(response_text)
            return json.loads(cleaned_text)
        except json.JSONDecodeError:
            pass

        # Step 3: Handle multiple JSON objects
        merged_data = {}
        
        # More robust approach: split by lines and reconstruct JSON objects
        lines = response_text.split('\n')
        current_json = []
        brace_count = 0
        
        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
                
            # Count braces in this line
            brace_count += stripped_line.count('{') - stripped_line.count('}')
            current_json.append(line)
            
            # When brace count returns to 0, we have a complete JSON object
            if brace_count == 0 and current_json:
                json_text = '\n'.join(current_json)
                
                try:
                    cleaned_json = self._sanitize_json_text(json_text)
                    parsed = json.loads(json_text)
                    merged_data.update(parsed)
                    print(f"[DEBUG] Successfully parsed JSON object with keys: {list(parsed.keys())}")
                except json.JSONDecodeError as e:
                    print(f"[JSON Parse Error] {e}")
                    print(f"[DEBUG] Problematic JSON text: {json_text[:200]}...")
                current_json = []
        
        return merged_data

    def _sanitize_json_text(self,text: str) -> str:
        """Sanitize text to make it valid JSON"""
        # Remove control characters
        sanitized = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
        
        # Fix common quote issues
        sanitized = sanitized.replace('"', '"').replace('"', '"')
        sanitized = sanitized.replace("'", "'").replace("�", "?")
        
        # Remove any trailing commas before closing braces/brackets
        sanitized = re.sub(r',(\s*[}\]])', r'\1', sanitized)
        
        # Ensure proper string escaping for newlines and quotes
        # This regex finds content within double quotes and escapes newlines
        def escape_newlines_in_strings(match):
            content = match.group(1)
            # Escape newlines and backslashes within the string
            content = content.replace('\\', '\\\\').replace('\n', '\\n').replace('\r', '\\r')
            return f'"{content}"'
        
        sanitized = re.sub(r'"([^"]*)"', escape_newlines_in_strings, sanitized)
        
        return sanitized


    def _post_process_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Enhanced post-processing with OCR-specific corrections"""
        # Ensure all required keys exist
        required_keys = ["personal_info", "skills", "experience", "education", "certifications", "projects", "languages", "awards"]
        for key in required_keys:
            if key not in data or data[key] is None:
                if key == "personal_info":
                    data[key] = {}
                else:
                    data[key] = []
        
        # Enhanced personal info cleaning
        if "personal_info" in data and data["personal_info"]:
            personal_info = data["personal_info"]
            
            # Fix common OCR errors in names
            if "name" in personal_info and personal_info["name"]:
                name = personal_info["name"]
                # Common OCR fixes for names
                name = re.sub(r'\bl\b', 'I', name)  # standalone 'l' to 'I'
                name = re.sub(r'\b0\b', 'O', name)  # standalone '0' to 'O'
                name = re.sub(r'[^\w\s\-\.]', '', name)  # Remove special chars except hyphens and dots
                name = ' '.join(word.capitalize() for word in name.split())  # Proper case
                personal_info["name"] = name.strip()
            
            # Enhanced email validation and correction
            if "email" in personal_info and personal_info["email"]:
                email = personal_info["email"]
                # Common OCR email fixes
                email = email.replace(' ', '')  # Remove spaces
                email = email.replace('(at)', '@')  # Common OCR substitution
                email = email.replace('[at]', '@')
                email = email.replace('(dot)', '.')
                email = email.replace('[dot]', '.')
                
                # Validate email format
                if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
                    personal_info["email"] = email.lower()
                else:
                    # Try to extract valid email from the text
                    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', email)
                    if email_match:
                        personal_info["email"] = email_match.group(0).lower()
                    else:
                        personal_info["email"] = ""
            
            # Enhanced phone number cleaning using the new helper method
            if "phone" in personal_info and personal_info["phone"]:
                personal_info["phone"] = self._clean_phone_number(personal_info["phone"])
            
            # Clean LinkedIn URL
            if "linkedin" in personal_info and personal_info["linkedin"]:
                linkedin = personal_info["linkedin"]
                # Common OCR fixes
                linkedin = linkedin.replace(' ', '')
                linkedin = linkedin.replace('|inkedin', 'linkedin')
                linkedin = linkedin.replace('Iinkedin', 'linkedin')
                linkedin = linkedin.replace('linkedln', 'linkedin')
                
                if "linkedin.com" in linkedin.lower():
                    personal_info["linkedin"] = linkedin
                elif linkedin.startswith("linkedin.com") or linkedin.startswith("www.linkedin.com"):
                    personal_info["linkedin"] = "https://" + linkedin
                else:
                    # Try to extract LinkedIn username
                    linkedin_match = re.search(r'linkedin\.com/in/([^/\s]+)', linkedin)
                    if linkedin_match:
                        personal_info["linkedin"] = f"https://linkedin.com/in/{linkedin_match.group(1)}"
                    else:
                        personal_info["linkedin"] = ""
        
        # Enhanced skills cleaning
        if "skills" in data and isinstance(data["skills"], list):
            skills = []
            seen_skills = set()
            for skill in data["skills"]:
                if skill and skill.strip():
                    # Clean skill name
                    skill_cleaned = re.sub(r'[^\w\s\+\#\-\.]', '', skill.strip())
                    skill_cleaned = ' '.join(word.capitalize() for word in skill_cleaned.split())
                    
                    # Skip if too short or too long (likely OCR errors)
                    if 2 <= len(skill_cleaned) <= 50 and skill_cleaned.lower() not in seen_skills:
                        skills.append(skill_cleaned)
                        seen_skills.add(skill_cleaned.lower())
            data["skills"] = skills
        
        # Enhanced experience cleaning
        if "experience" in data and isinstance(data["experience"], list):
            cleaned_experience = []
            for exp in data["experience"]:
                if isinstance(exp, dict):
                    # Clean company names
                    if "company" in exp and exp["company"]:
                        company = exp["company"]
                        company = re.sub(r'[^\w\s\-\&\.]', '', company)
                        company = ' '.join(word.capitalize() for word in company.split())
                        exp["company"] = company.strip()
                    
                    # Clean job titles
                    if "title" in exp and exp["title"]:
                        title = exp["title"]
                        title = re.sub(r'[^\w\s\-\&\.]', '', title)
                        title = ' '.join(word.capitalize() for word in title.split())
                        exp["title"] = title.strip()
                    
                    # Only keep if has meaningful content
                    if exp.get("company") and exp.get("title"):
                        cleaned_experience.append(exp)
            data["experience"] = cleaned_experience
        
        return data

    def get_usage_summary(self) -> Dict:
        """Get usage summary for all providers and models."""
        today = self._get_today_key()
        summary = {}
        
        for provider, config in self.api_providers.items():
            if not config['api_key']:
                continue
                
            provider_summary = {}
            total_usage = self._get_today_usage(provider)
            daily_limit = self.daily_limits.get(provider, 1000)
            
            provider_summary['total_usage'] = total_usage
            provider_summary['daily_limit'] = daily_limit
            provider_summary['remaining'] = daily_limit - total_usage
            provider_summary['usage_percentage'] = (total_usage / daily_limit) * 100
            provider_summary['status'] = 'Available' if total_usage < daily_limit else 'Exhausted'
            provider_summary['vision_models'] = {}
            provider_summary['text_models'] = {}
            
            # Vision model-specific usage
            for model_info in config['vision_models']:
                model = model_info[0]
                model_usage = self._get_today_usage(provider, model)
                provider_summary['vision_models'][model] = {
                    'usage': model_usage,
                    'last_used': 'Today' if model_usage > 0 else 'Not used today'
                }
            
            # Text model-specific usage
            for model_info in config['text_models']:
                model = model_info[0]
                model_usage = self._get_today_usage(provider, model)
                provider_summary['text_models'][model] = {
                    'usage': model_usage,
                    'last_used': 'Today' if model_usage > 0 else 'Not used today'
                }
            
            summary[provider] = provider_summary
        
        return summary

    def extract_batch(self, file_paths: List[str], output_file: str = None) -> List[Dict[str, Any]]:
        """Extract data from multiple resume files with better progress tracking"""
        results = []
        total_files = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            filename = os.path.basename(file_path)
            logger.info(f"Processing {i+1}/{total_files}: {filename}")
            
            result = self.extract_from_file(file_path, filename)
            results.append(result)
            
            # Show progress
            if result["success"]:
                logger.info(f"✓ Successfully extracted data from {filename}")
                if "personal_info" in result and result["personal_info"].get("name"):
                    logger.info(f"  Candidate: {result['personal_info']['name']}")
            else:
                logger.warning(f"✗ Failed to extract data from {filename}: {result.get('error', 'Unknown error')}")
            
            # Small delay to avoid overwhelming the API
            time.sleep(0.1)
        
        # Save results if output file specified
        if output_file:
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                logger.info(f"Results saved to: {output_file}")
            except Exception as e:
                logger.error(f"Error saving results: {e}")
        
        # Print summary
        successful = sum(1 for r in results if r["success"])
        logger.info(f"Batch processing complete: {successful}/{total_files} files processed successfully")
        
        return results

    async def extract_batch_async(self, file_paths: List[str], progress_callback=None) -> List[Dict[str, Any]]:
        """Async version of batch extraction for better integration with FastAPI"""
        results = []
        total_files = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            filename = os.path.basename(file_path)
            
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self.extract_from_file, file_path, filename)
            results.append(result)
            
            # Call progress callback if provided
            if progress_callback:
                progress = int((i + 1) * 100 / total_files)
                await progress_callback(progress)
            
            # Small delay to avoid overwhelming the API
            await asyncio.sleep(0.1)
        
        return results


if __name__ == "__main__":
    import argparse
    import os
    import json
    import logging
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='Resume Extractor - Manual PDF Test with Google Vision')
    parser.add_argument('pdf_path', help='Path to the PDF file to analyze')
    parser.add_argument('--output', '-o', help='Output JSON file path (optional)')
    args = parser.parse_args()
    
    # Check if file exists
    if not os.path.exists(args.pdf_path):
        print(f"Error: File not found: {args.pdf_path}")
        exit(1)
    
    # Initialize the extractor
    try:
        extractor = ResumeExtractor()
        print(f"Initializing ResumeExtractor... Done.")
        
        # Print OCR capabilities
        ocr_caps = extractor.get_ocr_capabilities()
        print("OCR Capabilities:")
        for method, available in ocr_caps.items():
            print(f"  - {method}: {'Available' if available else 'Not available'}")
        
        # Extract resume data directly using Google Vision
        filename = os.path.basename(args.pdf_path)
        print(f"Processing PDF with Google Vision OCR: {filename}")
        result = extractor.extract_from_file(args.pdf_path, filename)
        
        # Print extraction results
        print("\nExtraction Results:")
        print(f"Success: {result.get('success', False)}")
        
        if result.get('success', False):
            print(f"Name: {result.get('personal_info', {}).get('name', 'N/A')}")
            print(f"Email: {result.get('personal_info', {}).get('email', 'N/A')}")
            print(f"Phone: {result.get('personal_info', {}).get('phone', 'N/A')}")
            print(f"Skills: {len(result.get('skills', []))} found")
            print(f"Experience: {len(result.get('experience', []))} entries")
            print(f"Education: {len(result.get('education', []))} entries")
            
            # Save to output file if specified
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2)
                print(f"\nResults saved to: {args.output}")
            else:
                print("\nFull JSON result:")
                print(json.dumps(result, indent=2))
        else:
            print(f"Error: {result.get('error', 'Unknown error')}")
            if 'raw_response' in result:
                print(f"Raw response: {result['raw_response']}")
    
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()